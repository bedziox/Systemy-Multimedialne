import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from tqdm import tqdm
import sys
from docx import Document
from docx.shared import Inches


def imgToUInt8(img):
    if np.issubdtype(img.dtype,np.unsignedinteger):
        return img
    else:
        scaledImg = img * 255
        scaledImg = pd.DataFrame.astype(scaledImg, type='uint32')
        return scaledImg
    
# data = 
# przypadki testowe
samples = [
    np.array([1,1,1,1,2,1,1,1,1,2,1,1,1,1]),
    np.array([1,2,3,1,2,3,1,2,3]),
    np.array([5,1,5,1,5,5,1,1,5,5,1,1,5]),
    np.array([-1,-1,-1,-5,-5,-3,-4,-2,1,2,2,1]),
    np.zeros((1,520)),
    np.arange(0,521,1),
    np.eye(7),
    np.dstack([np.eye(7),np.eye(7),np.eye(7)]),
    np.ones((1,1,1,1,1,1,10))
    ]

def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

# ========== RLE ========== #
def rleEncoder(img):
    encoded = np.zeros(np.prod(img.shape)*2)
    img1d = img.flatten()
    info = np.array([len(img.shape)])
    info = np.concatenate([info, img.shape])
    i = 0
    bufferPosition = 0
    while i < len(img1d):
        counter = compareNext(img1d[i:len(img1d)])
        encoded[bufferPosition] = counter
        encoded[bufferPosition+1] =  img1d[i]
        bufferPosition += 2
        i += counter
    encoded = encoded[:bufferPosition]
    encoded = np.concatenate([info, encoded])
    return encoded.astype(int)

def rleDecoder(encoded):
    shape = encoded[1:int(encoded[0]+1)]
    decoded = []
    i = encoded[0] + 1
    while i < len(encoded):
        decoded = pushNumbers(decoded, encoded[i], encoded[i+1])    
        i += 2
    decoded = decoded.reshape(shape)
    return decoded

# ========== Byte run ========== #  

def byteRunEncoder(img):
    encoded = np.zeros(np.prod(img.shape)*2)
    img1d = img.flatten()
    info = np.concatenate([np.array([len(img.shape)]), img.shape])
    i = 0
    bufferPosition = 0
    while i < len(img1d):
        counter = compareNextByteRun(img1d[i:len(img1d)])
        if counter > 1:
            encoded[bufferPosition] = -counter + 1
            encoded[bufferPosition+1] = img1d[i]
            bufferPosition += 2
            i += counter 
        else:
            num = countDifferent(img1d[i:len(img1d)])
            encoded[bufferPosition] = num - 1
            bufferPosition += 1
            for j in range (num):
                encoded[bufferPosition + j] = img1d[i + j]
            bufferPosition += num
            i += num
    encoded = encoded[:bufferPosition]
    encoded = np.concatenate([info, encoded])
    return encoded.astype(int)

def byteRunDecoder(encoded):
    shape = encoded[1:encoded[0]+1]
    decoded = np.zeros(np.prod(shape)*2)
    i = int(encoded[0] + 1)
    bufferPosition = 0
    while i < len(encoded) - 1:
        if encoded[i] < 0:
            countToPush = -encoded[i] + 1
            valueToPush = encoded[i + 1]
            for _ in range(countToPush):
                decoded[bufferPosition] = valueToPush
                bufferPosition += 1
            i += 2
        else:
            numbers = encoded[i] + 1
            i += 1
            for _ in range(numbers):
                decoded[bufferPosition] = encoded[i] 
                i += 1
                bufferPosition += 1
    decoded = decoded[:bufferPosition]
    decoded = decoded.reshape(shape.astype(int))
    return decoded

# ========== UTIL ========== #        
def compareNext(data):
    counter = 1
    for j in range(0, len(data) - 1):
        if(data[j] == data[j+1]):
            counter = counter + 1
        else:
            break
    return counter

def compareNextByteRun(data):
    counter = 1
    for j in range(0, len(data) - 1):
        if(data[j] == data[j+1]):
            counter = counter + 1
            if counter == 127:
                return counter
        else:
            break
    return counter

def countDifferent(data):
    counter = 1
    for j in range(0, len(data) - 1):
        if(data[j] != data[j+1]):
            counter = counter + 1
            if counter == 127:
                return counter
        else:
            break
    return counter

def pushNumbers(array, count, number):
    array = np.array(array)
    array.resize(len(array) + int(count))
    array[-int(count):] = number
    return array

document = Document()
document.add_heading('Lab 5',0)

# ========== EXECUTION - Test data ========== #
document.add_heading("Dane testowe:", 2)
for index, data in enumerate(samples):
    document.add_paragraph(f"dane testowe nr: {index + 1} \n{data}")
    encodedDataRLE = rleEncoder(data)
    decodedDataRLE = rleDecoder(encodedDataRLE)
    
    encodedDataBR = byteRunEncoder(data)
    decodedDataBR = byteRunDecoder(encodedDataBR)
    document.add_paragraph(f" RLE: {decodedDataRLE.all() == data.all()}")
    document.add_paragraph(f" ByteRun: {decodedDataBR.all() == data.all()}")

document.add_heading("Pliki:", 2)
files = ["imgs/wzor_dokumentu.jpg", "imgs/rysunek_techniczny.jpg", "imgs/kapibara.jpeg"]

for file in files:
    print(f"{file} start")
    img = plt.imread(file)
    img = img.astype(int)
    document.add_heading(f"{file}", 3)
    document.add_picture(file)
    size = get_size(img)
        
    print(f"{file} rle encoding")
    encodedImgRLE = rleEncoder(img)
    lengthEncodedRLE = get_size(encodedImgRLE)
    print(f"{file} rle decoding")
    decodedImgRLE = rleDecoder(encodedImgRLE)
    CRRLE = size/lengthEncodedRLE
    PRRLE = lengthEncodedRLE*100/size
    document.add_paragraph(f" RLE: {decodedImgRLE.all() == img.all()}")
    document.add_paragraph(f"CR RLE:{CRRLE} PR RLE:{PRRLE}")
    
    print(f"{file} byterun encoding")
    encodedImgBT = byteRunEncoder(img)
    lengthEncodedBT = get_size(encodedImgBT)
    print(f"{file} byterun decoding")
    decodedImgBT = byteRunDecoder(encodedImgBT)
    CRBT = size/lengthEncodedBT
    PRBT = lengthEncodedBT*100/size
    document.add_paragraph(f" ByteRun: {decodedImgBT.all() == img.all()}")
    document.add_paragraph(f"CR ByteRun:{CRBT} PR ByteRun:{PRBT}")
    print(f"{file} finish")

document.save("lab5.docx")