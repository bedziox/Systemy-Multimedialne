import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
import cv2
from docx import Document
from docx.shared import Inches
from io import BytesIO

img1 = plt.imread("IMG_INTRO/A1.png")

print(img1.dtype)
print(img1.shape)
print(np.min(img1),np.max(img1))

img2 = plt.imread("IMG_INTRO/A2.jpg")

print(img2.dtype)
print(img2.shape)
print(np.min(img2),np.max(img2))

img3 = cv2.imread("IMG_INTRO/A3.png")

print(img3.dtype)
print(img3.shape)
print(np.min(img3),np.max(img3))

img3 = cv2.imread("IMG_INTRO/A4.jpg")

print(img3.dtype)
print(img3.shape)
print(np.min(img3),np.max(img3))

def imgToUInt8(img):
    if np.issubdtype(img.dtype,np.unsignedinteger):
        return img
    else:
        scaledImg = img * 255
        scaledImg = pd.DataFrame.astype(scaledImg, type='uint8')
        return scaledImg
        

def imgToFloat(img):
    if np.issubdtype(img.dtype,np.floating):
        return img
    else:
        scaledImg = img / 255
        return scaledImg

plt.imshow(img3)
plt.show()

R = img3[:,:,2]
G = img3[:,:,1]
B = img3[:,:,0]

plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=255)
plt.show()

Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

plt.imshow(Y2, cmap=plt.cm.gray )
plt.show()

# na przyszlosc zmiana formatu kolorow
# img_RGB = cv2.cvtColor(img_BGR, cv2.COLOR_BGR2RGB)
# img_BGR = cv2.cvtColor(img_RGB, cv2.COLOR_RGB2BGR)

#zadanie 2

#img 1

imgB1 = plt.imread("IMG_INTRO/B01.png")
print(np.min(imgB1),np.max(imgB1))

R = imgB1[:,:,0]
G = imgB1[:,:,1]
B = imgB1[:,:,2]
Y1 = 0.229 * R + 0.587 * G + 0.114 * B
Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

plt.subplot(3,3,1)
plt.imshow(imgB1)
plt.title("Oryginalny")

plt.subplot(3,3,2)
plt.imshow(Y1, cmap=plt.cm.gray)
plt.title("Y1")

plt.subplot(3,3,3)
plt.imshow(Y2, cmap=plt.cm.gray)
plt.title("Y2")

plt.subplot(3,3,4)
plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=1)
plt.title("R")

plt.subplot(3,3,5)
plt.imshow(G, cmap=plt.cm.gray, vmin=0, vmax=1)
plt.title("G")

plt.subplot(3,3,6)
plt.imshow(B, cmap=plt.cm.gray, vmin=0, vmax=1)
plt.title("B")

imgB1copy = imgB1.copy()
imgB1copy[:,:,1] = 0
imgB1copy[:,:,2] = 0

plt.subplot(3,3,7)
plt.imshow(imgB1copy)
plt.title("Just R")

imgB1copy2 = imgB1.copy()
imgB1copy2[:,:,0] = 0
imgB1copy2[:,:,2] = 0

plt.subplot(3,3,8)
plt.imshow(imgB1copy2)
plt.title("Just B")

imgB1copy2 = imgB1.copy()
imgB1copy2[:,:,0] = 0
imgB1copy2[:,:,1] = 0

plt.subplot(3,3,9)
plt.imshow(imgB1copy2)
plt.title("Just G")

plt.show()

# image 2

def fragmentPlots(partImg, index, filename = 0):

    R = partImg[:,:,0]
    G = partImg[:,:,1]
    B = partImg[:,:,2]

    Y1 = 0.229 * R + 0.587 * G + 0.114 * B
    Y2 = 0.2126 * R + 0.7152 * G + 0.0722 * B

    plt.subplot(3,3,1)

    plt.imshow(partImg)
    plt.title("Oryginalny")

    plt.subplot(3,3,2)
    plt.imshow(Y1, cmap=plt.cm.gray)
    plt.title("Y1")

    plt.subplot(3,3,3)
    plt.imshow(Y2, cmap=plt.cm.gray)
    plt.title("Y2")

    plt.subplot(3,3,4)
    plt.imshow(R, cmap=plt.cm.gray, vmin=0, vmax=255)
    plt.title("R")

    plt.subplot(3,3,5)
    plt.imshow(G, cmap=plt.cm.gray, vmin=0, vmax=255)
    plt.title("G")

    plt.subplot(3,3,6)
    plt.imshow(B, cmap=plt.cm.gray, vmin=0, vmax=255)
    plt.title("B")

    imgCopy = partImg.copy()
    imgCopy[:,:,1] = 0
    imgCopy[:,:,2] = 0

    plt.subplot(3,3,7)
    plt.imshow(imgCopy)
    plt.title("Just R")

    imgCopy2 = partImg.copy()
    imgCopy2[:,:,0] = 0
    imgCopy2[:,:,2] = 0

    plt.subplot(3,3,8)
    plt.imshow(imgCopy2)
    plt.title("Just G")

    imgCopy3 = partImg.copy()
    imgCopy3[:,:,0] = 0
    imgCopy3[:,:,1] = 0

    plt.subplot(3,3,9)
    plt.imshow(imgCopy3)
    plt.title("Just B")
    
    plt.savefig("Image{}".format(index))
    plt.close()
    return "Image{}".format(index)
    

imgB2 = plt.imread("IMG_INTRO/B02.jpg")
imgB2 = imgToUInt8(imgB2)
fragment = imgB2[0:400,0:600].copy()
fragmentPlots(fragment,"1")


df = pd.DataFrame(data={'Filename':['IMG_INTRO/B02.jpg'],'Grayscale':[False],
                        'Fragments': [
                            [[0, 0, 200, 200], [200, 200, 400, 400], [400, 400, 600, 600], [600, 600, 800, 800],
                             [500, 500, 700, 700]]]
                        })

            
            
document = Document()
document.add_heading('Report',0) # tworzenie nagłówków druga wartość to poziom nagłówka 
memfile = BytesIO()



for index, row in df.iterrows():
    img = plt.imread(row['Filename'])
    counter = 0
    if row['Fragments'] is not None:
        for f in row['Fragments']:
            fragment = img[f[0]:f[2],f[1]:f[3]].copy()
            # tu wykonujesz operacje i inne wyświetlenia na fragmencie
            file = fragmentPlots(fragment,counter,row['Filename'])
            document.add_heading(f"Plik nr {counter}",2) # nagłówek sekcji, mozę być poziom wyżej
            document.add_picture(f'{file}.png', width=Inches(6)) # dodanie obrazu z bufora do pliku
            counter = counter+1

document.save('lab2report.docx') # zapis do pliku