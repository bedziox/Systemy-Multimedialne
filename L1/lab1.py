import numpy as np
import matplotlib.pyplot as plt
import scipy.fftpack
import sounddevice as sd
import soundfile as sf

from docx import Document
from docx.shared import Inches
from io import BytesIO


# Zadanie 1
data, fs = sf.read('SOUND_INTRO/sound1.wav', dtype='float32') 
print(fs)

print(data.dtype)
print(data.shape)

sd.play(data, fs)
status = sd.wait()

sf.write('sound_L.wav', data[:,0], fs)
sf.write('sound_R.wav', data[:,1], fs)
sf.write('sound_mix.wav', np.mean(data, axis=1), fs)

t = data.shape[0]/fs
x = np.linspace(0,t,data.shape[0])

plt.subplot(2,1,1)
plt.plot(x,data[:,0])

plt.subplot(2,1,2)
plt.plot(x,data[:,1])
plt.show()


data, fs = sf.read('SIN/sin_440Hz.wav', dtype=np.int32)

fsize=2**8

plt.figure()
plt.subplot(2,1,1)
plt.plot(np.arange(0,data.shape[0])/fs,data)

plt.subplot(2,1,2)
yf = scipy.fftpack.fft(data,fsize)
plt.plot(np.arange(0,fs/2,fs/fsize),20*np.log10( np.abs(yf[:fsize//2])))
plt.show()

# Zadanie 2
def plotAudio(axs, Signal,Fs,TimeMargin=[0,0.02], fsize = 2**8):
    
    axs[0].plot(np.arange(0,Signal.shape[0])/Fs,Signal)
    axs[0].set(ylabel= "A", xlabel = "T [s]")
    axs[0].set(xlim = TimeMargin)

    yf = scipy.fftpack.fft(Signal,fsize)
    x = np.arange(0,fs/2,fs/fsize)
    y = 20*np.log10( np.abs(yf[:fsize//2]))
    axs[1].plot(x,y)
    axs[1].set(ylabel= "dB", xlabel = "F [Hz]")
    maxValue = [x[np.argmax(y)], y[np.argmax(y)]]
    
    return maxValue
plotAudio(data, fs)




document = Document()
document.add_heading('Lab1 Zad3',0) # tworzenie nagłówków druga wartość to poziom nagłówka 

files=['SIN/sin_60Hz.wav','SIN/sin_440Hz.wav','SIN/sin_8000Hz.wav']
fsizes=[2**8,2**12,2**16]
for file in files:
    document.add_heading('Plik - {}'.format(file),2)
    for i,fsize in enumerate(fsizes):
        document.add_heading('Fsize {}'.format(fsize),3) # nagłówek sekcji, mozę być poziom wyżej
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
    
        ############################################################
        # Tu wykonujesz jakieś funkcje i rysujesz wykresy
        ############################################################
        data, fs = sf.read(file, dtype=np.int32)
        maxValue = plotAudio(axs, data, fs, fsize=fsize)
        
        fig.suptitle('Fsize {}'.format(fsize)) # Tytuł wykresu
        fig.tight_layout(pad=1.5) # poprawa czytelności 
        memfile = BytesIO() # tworzenie bufora
        fig.savefig(memfile) # z zapis do bufora 
        
    
        document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
        
        memfile.close()
        ############################################################
        # Tu dodajesz dane tekstowe - wartosci, wyjscie funkcji ect.
        document.add_paragraph('wartość maksymalna = {}'.format(maxValue[0])) 
        document.add_paragraph('wartość w punkcie maksymalnym = {}'.format(maxValue[1])) 
        ############################################################

document.save('report.docx') # zapis do pliku