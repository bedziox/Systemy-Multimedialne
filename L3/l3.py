import numpy as np
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os


def colorFit(pixel, Pallet):
    return Pallet[np.argmin(np.linalg.norm(Pallet-pixel,axis=1))]

pallet8 = np.array([
        [0.0, 0.0, 0.0,],
        [0.0, 0.0, 1.0,],
        [0.0, 1.0, 0.0,],
        [0.0, 1.0, 1.0,],
        [1.0, 0.0, 0.0,],
        [1.0, 0.0, 1.0,],
        [1.0, 1.0, 0.0,],
        [1.0, 1.0, 1.0,],
])

pallet16 =  np.array([
        [0.0, 0.0, 0.0,], 
        [0.0, 1.0, 1.0,],
        [0.0, 0.0, 1.0,],
        [1.0, 0.0, 1.0,],
        [0.0, 0.5, 0.0,], 
        [0.5, 0.5, 0.5,],
        [0.0, 1.0, 0.0,],
        [0.5, 0.0, 0.0,],
        [0.0, 0.0, 0.5,],
        [0.5, 0.5, 0.0,],
        [0.5, 0.0, 0.5,],
        [1.0, 0.0, 0.0,],
        [0.75, 0.75, 0.75,],
        [0.0, 0.5, 0.5,],
        [1.0, 1.0, 1.0,], 
        [1.0, 1.0, 0.0,]
])

# print(colorFit(np.array([0.25,0.25,0.5]),pallet8))
# print(colorFit(np.array([0.25,0.25,0.5]),pallet16))

def kwant_colorFit(img,Pallet):
        out_img = img.copy()
        for w in range(out_img.shape[0]):
            for k in range(out_img.shape[1]):
                out_img[w,k]=colorFit(img[w,k],Pallet)
        return out_img

def imgToFloat(img):
    if np.issubdtype(img.dtype,np.floating):
        return img
    else:
        scaledImg = img / 255
        return scaledImg



img1 = imgToFloat(plt.imread("IMG_GS/GS_0001.tif"))
img2 = plt.imread("IMG_GS/GS_0002.png")
img3 = plt.imread("IMG_GS/GS_0003.png")

def plotKwantImgs(img, index):
    plt.subplot(1,4,1)
    plt.imshow(img, cmap=plt.cm.gray)
    plt.title("Oryginalne")
    
    plt.subplot(1,4,2)
    N = 2**1
    img1bit = img.copy()
    paleta = np.linspace(0,1,N).reshape(N,1)
    plt.imshow(kwant_colorFit(img1bit, paleta), cmap=plt.cm.gray)
    plt.title("Paleta 1 bit")
    
    plt.subplot(1,4,3)
    N = 2**2
    img2bit = img.copy()
    paleta = np.linspace(0,1,N).reshape(N,1)
    plt.imshow(kwant_colorFit(img2bit, paleta), cmap=plt.cm.gray)
    plt.title("Paleta 2 bit")
    
    plt.subplot(1,4,4)
    N = 2**4
    img4bit = img.copy()
    paleta = np.linspace(0,1,N).reshape(N,1)
    plt.imshow(kwant_colorFit(img4bit, paleta), cmap=plt.cm.gray)
    plt.title("Paleta 4 bit")
    
    plt.savefig("RESULTS/Kwant{}".format(index))
    
def plotDither1Bit(img, paleta):
    plt.title("Dithering - Pallet 1 bit") 
    plt.subplot(2,3,1)
    plt.imshow(img, cmap=plt.cm.gray)
    plt.title("Oryginalne")
    
    plt.subplot(2,3,2)
    img1 = img.copy()
    plt.imshow(kwant_colorFit(img1, paleta), cmap=plt.cm.gray)
    plt.title("Kwantyzacja")
    
    plt.subplot(2,3,3)
    img2 = img.copy()
    plt.imshow(randomDithering(img2), cmap=plt.cm.gray)
    plt.title("Dithering losowy")
    
    plt.subplot(2,3,4)
    img3 = img.copy()
    plt.imshow(orderDithering(img3, M2, paleta), cmap=plt.cm.gray)
    plt.title("Dithering ord")
    
    plt.subplot(2,3,5)
    img4 = img.copy()
    plt.imshow(fsDithering(img4, paleta), cmap=plt.cm.gray)
    plt.title("Dithering FS")
    
    plt.savefig("RESULTS/Dither1Bit{}".format(index))

def plotDitherGreyscale(img, paleta, name):
    plt.title(name)
    plt.subplot(2,2,1)
    plt.imshow(img, cmap=plt.cm.gray)
    plt.title("Oryginalne")
    
    plt.subplot(2,2,2)
    img1 = img.copy()
    plt.imshow(kwant_colorFit(img1, paleta), cmap=plt.cm.gray)
    plt.title("Kwantyzacja")
    
    plt.subplot(2,2,3)
    img2 = img.copy()
    plt.imshow(orderDithering(img2, M2, paleta), cmap=plt.cm.gray)
    plt.title("Dithering ord")
    
    plt.subplot(2,2,4)
    img3 = img.copy()
    plt.imshow(fsDithering(img3, paleta), cmap=plt.cm.gray)
    plt.title("Dithering FS")
    
    plt.savefig("RESULTS/Dither{}.png".format(name))
    
def plotDither(img, paleta, name):
    plt.title(name)
    plt.subplot(2,2,1)
    plt.imshow(img, cmap=plt.cm.gray)
    plt.title("Oryginalne")
    
    plt.subplot(2,2,2)
    img1 = img.copy()
    plt.imshow(kwant_colorFit(img1, paleta))
    plt.title("Kwantyzacja")
    
    plt.subplot(2,2,3)
    img2 = img.copy()
    plt.imshow(orderDithering(img2, M2, paleta))
    plt.title("Dithering ord")
    
    plt.subplot(2,2,4)
    img3 = img.copy()
    plt.imshow(fsDithering(img3, paleta))
    plt.title("Dithering FS")
    
    plt.savefig("RESULTS/Dither{}.png".format(name))

    
    
# plotKwantImgs(img1, 1)
# plotKwantImgs(img2[:,:,0], 2)
# plotKwantImgs(img3[:,:,0], 3)


def randomDithering(img):
    return (img.copy()>=np.random.rand(img.shape[0],img.shape[1])) * 1

# img1dit = randomDithering(img1)
# plt.figure()
# plt.imshow(img1dit, cmap=plt.cm.gray)
# plt.show()


M2 = np.array([
    [0, 8, 2, 10],
    [12, 4, 14, 6],
    [3, 11, 1, 9],
    [15, 7, 13, 5]])



def orderDithering(img, M, Pallet, r = 1):
    Mpre = (M+1) / (2*np.sqrt(M.shape[0]))**2 - 0.5
    out_img = img.copy()
    for w in range(out_img.shape[0]):
        for k in range(out_img.shape[1]):
            out_img[w,k]=colorFit((img[w,k] + r * Mpre[w % M.shape[0], k % M.shape[0]]), Pallet)
    return out_img

def fsDithering(img, Pallet):
    out_img = img.copy()
    for w in range(out_img.shape[0]-1):
        for k in range(out_img.shape[1]-1):
            oldpixel = out_img[w,k]
            newpixel = colorFit(oldpixel, Pallet)
            out_img[w,k] = newpixel
            quant_error = oldpixel - newpixel
            if w + 1 < img.shape[0]:
                out_img[w + 1, k] = out_img[w + 1, k] + quant_error * 7 / 16
            if k + 1 < img.shape[1]:
                out_img[w - 1, k + 1] = out_img[w - 1, k + 1] + quant_error * 3 / 16
                if w > 0:
                    out_img[w, k + 1] = out_img[w, k + 1] + quant_error * 5 / 16
                if w + 1 < img.shape[0]:
                    out_img[w + 1, k + 1] = out_img[w + 1, k + 1] + quant_error * 1 / 16
    return out_img

N = 2 ** 1
paleta = np.linspace(0,1,N).reshape(N,1)
img1dit = plt.imread("IMG_SMALL/SMALL_0002.png")

print(np.unique(fsDithering(img1dit.copy(),np.linspace(0,1,2).reshape(2,1))).size)

# df = pd.DataFrame(data={'Filename':['IMG_SMALL/SMALL_0001.tif','IMG_SMALL/SMALL_0002.png','IMG_SMALL/SMALL_0003.png', 'IMG_SMALL/SMALL_0004.jpg', 'IMG_SMALL/SMALL_0006.jpg', 'IMG_SMALL/SMALL_0007.jpg', 'IMG_SMALL/SMALL_0009.jpg'],'Grayscale':[True, True, True, False, False, False, False]
#                         })       
# document = Document()
# document.add_heading('Report',0)

# for index, row in df.iterrows():
#     img = plt.imread(row['Filename'])
#     img = imgToFloat(img)
#     counter = 0
#     if row['Grayscale']:
#         img = img[:,:,1]
#         plotDither1Bit(img, np.linspace(0,1,2**1).reshape(2**1,1))
#         plotDitherGreyscale(img, np.linspace(0,1,2**2).reshape(2**2,1), "greyscale2bit{}".format(index))
#         plotDitherGreyscale(img, np.linspace(0,1,2**4).reshape(2**4,1), "greyscale4bit{}".format(index))
#     else:
#         img = img[:,:,:3]
#         plotDither(img, pallet8, "pallet8{}".format(index))
#         plotDither(img, pallet16, "pallet16{}".format(index))

# for filename in os.listdir("RESULTS"):
#     img_path = os.path.join("RESULTS",filename)
#     name, sep, extension = filename.partition('.')
#     document.add_heading(name, 2)
#     document.add_picture(img_path, width=Inches(6))

# document.save('lab3report.docx') # zapis do pliku