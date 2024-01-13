from docx import Document
from docx.shared import Inches

import numpy as np
import matplotlib.pyplot as plt
import sounddevice as sd
import soundfile as sf
import scipy.fftpack

from scipy.interpolate import interp1d
import matplotlib.pyplot as plt

from io import BytesIO

dataSing1, fsSing1 = sf.read('SING/sing_low1.wav', dtype='float32') 
dataSing2, fsSing2 = sf.read('SING/sing_medium1.wav', dtype='float32')  
dataSing3, fsSing3 = sf.read('SING/sing_high1.wav', dtype='float32') 


Sing1 = ["sing_low1",dataSing1, fsSing1]
Sing2 = ["sing_medium1",dataSing2, fsSing2]
Sing3 = ["sing_high1",dataSing3, fsSing3]

files = [Sing1, Sing2, Sing3]

# np.sign()
# np.log()
# np.exp()

def Kwant(data,bit):
    d= 2 ** bit - 1
    if np.issubdtype(data.dtype,np.floating):
        m = -1
        n = 1
    else:
        n = np.iinfo(data.dtype).min
        m = np.iinfo(data.dtype).max
    DataF = data.astype(float)
    DataF = (DataF - m) / (n - m) * d
    DataF = np.round(DataF)
    DataF = DataF / d * (n - m) + m
    # kwantyzacja na DataF
    return DataF.astype(data.dtype)

# LAWS

def a_law_encoder(data, A = 87.6):
    encoded = data.copy()
    idx = np.abs(data) < 1/A
    encoded[idx] = np.sign(data[idx]) * (A* np.abs(data[idx]) / (1 + np.log(A)))
    encoded[np.logical_not(idx)] = np.sign(data[np.logical_not(idx)]) * (1 + np.log(A * np.abs(data[np.logical_not(idx)]))) / (1 + np.log(A))
    return encoded 

def a_law_decoder(encoded, A = 87.6):
    decoded = np.copy(encoded)
    idx = np.abs(encoded) < 1/(1+np.log(A))
    decoded[idx] = np.sign(encoded[idx]) * (np.abs(encoded[idx]) * (1 + np.log(A))) / A
    decoded[np.logical_not(idx)] = np.sign(encoded[np.logical_not(idx)]) * np.exp(np.abs(encoded[np.logical_not(idx)]) * (1 + np.log(A)) - 1) / A
    return decoded

def u_law_encoder(data, u = 255):
    return np.sign(data) * np.log(1+u*np.abs(data)) / (np.log(1+u))


def u_law_decoder(encoded, u = 255):
    return np.sign(encoded) * 1/u * ((1+u)**np.abs(encoded) - 1)

# DPCM

def DPCM_compress(x,bit):
    y=np.zeros(x.shape)
    e=0
    for i in range(0,x.shape[0]):
        y[i]=Kwant(x[i]-e,bit)
        e+=y[i]
    return y

def DPCM_compress_pred(x,bit,predictor,n): 
    y=np.zeros(x.shape)
    xp=np.zeros(x.shape)
    e=0
    for i in range(1,x.shape[0]):
        y[i]=Kwant(x[i]-e,bit)
        xp[i]=y[i]+e
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        e=predictor(xp[idx])
    return y

def DPCM_decompress(y):
    x = np.zeros(y.shape)
    for i in range(0, x.shape[0] - 1):
        x[i] = y[i]
        y[i+1] += y[i]
    return x

def DPCM_decompress_pred(y, predictor, n):
    x = np.zeros(y.shape)
    e = 0
    for i in range(0, x.shape[0] - 1):
        x[i] = y[i]
        idx=(np.arange(i-n,i,1,dtype=int)+1)
        idx=np.delete(idx,idx<0)
        y[i+1] += e
        e=predictor(y[idx])
    return x

def no_pred(X):
    return X[-1]

def predictor(X):
    return np.mean(X[:3])

# TESTS

x=np.linspace(-1,1,1000)
y=0.9*np.sin(np.pi*x*4)

document = Document()
document.add_heading('Lab 6',0)

document.add_heading("Dane testowe, sinus", 1)

fig, ax = plt.subplots(5,1, figsize = (10,7), constrained_layout = True)
# fig.tight_layout(pad = 2)
fig.suptitle("Przyklad komprezji z kwantyzacja do 6 bitów")
ax[0].set_title("Sygnal oryginalny")
ax[0].plot(x, y)

ax[1].set_title("kwantyzacja 6 bit, a law")
aLawEncoded = a_law_encoder(y)
kwantyzacja = Kwant(aLawEncoded, 6)
aLawDecoded = a_law_decoder(kwantyzacja)
ax[1].plot(x,aLawDecoded)

ax[2].set_title("kwantyzacja 6 bit, mu law")
uLawEncoded = u_law_encoder(y)
kwantyzacja = Kwant(uLawEncoded, 6)
uLawDecoded = u_law_decoder(kwantyzacja)
ax[2].plot(x,uLawDecoded)

ax[3].set_title("kwantyzacja 6 bit, DPCM bez predykcji")
dpcmCompress = DPCM_compress(y, 6)
decmDecompress = DPCM_decompress(dpcmCompress)
ax[3].plot(x,decmDecompress)

ax[4].set_title("kwantyzacja 6 bit, DPCM z predykcją")
dpcmCompress = DPCM_compress_pred(y, 6, predictor=predictor, n = 3)
dpcmDecompress = DPCM_decompress_pred(dpcmCompress, predictor=predictor, n = 3)
ax[4].plot(x,dpcmCompress)
fig.savefig("testplt1")
document.add_picture("testplt1.png", width=Inches(6))

x=np.linspace(-0.5,-0.25,1000)
y=0.9*np.sin(np.pi*x*4)

fig, ax = plt.subplots(5,1, figsize = (10,7), constrained_layout = True)
# fig.tight_layout(pad = 2)
fig.suptitle("Przyklad komprezji z kwantyzacja do 6 bitów")
ax[0].set_title("Sygnal oryginalny")
ax[0].plot(x, y)

ax[1].set_title("kwantyzacja 6 bit, a law")
aLawEncoded = a_law_encoder(y)
kwantyzacja = Kwant(aLawEncoded, 6)
aLawDecoded = a_law_decoder(kwantyzacja)
ax[1].plot(x,aLawDecoded)

ax[2].set_title("kwantyzacja 6 bit, mu law")
uLawEncoded = u_law_encoder(y)
kwantyzacja = Kwant(uLawEncoded, 6)
uLawDecoded = u_law_decoder(kwantyzacja)
ax[2].plot(x,uLawDecoded)

ax[3].set_title("kwantyzacja 6 bit, DPCM bez predykcji")
dpcmCompress = DPCM_compress(y, 6)
decmDecompress = DPCM_decompress(dpcmCompress)
ax[3].plot(x,decmDecompress)

ax[4].set_title("kwantyzacja 6 bit, DPCM z predykcją")
dpcmCompress = DPCM_compress_pred(y, 6, predictor=predictor, n = 3)
dpcmDecompress = DPCM_decompress_pred(dpcmCompress, predictor=predictor, n = 3)
ax[4].plot(x,dpcmCompress)
fig.savefig("testplt2")

document.add_picture("testplt2.png", width=Inches(6))

document.add_heading("Pliki Sing", 1)

played = True
if not played:
    for file in files:
        document.add_heading(f"{file[0]}", 2)
        sd.play(file[1],file[2])
        sd.wait()
        input("Next song")
        for bit in [8,7,6,5,4,3,2]:
            print(f" {file[0]} kompresja a_law, kwantyzacja {bit} bit")
            aLawEncoded = a_law_encoder(file[1])
            kwantyzacja = Kwant(aLawEncoded, bit)
            aLawDecoded = a_law_decoder(kwantyzacja)
            sd.play(aLawDecoded,file[2])
            sd.wait()
            input("Next song")

            print(f" {file[0]} kompresja u_law, kwantyzacja {bit} bit")
            uLawEncoded = u_law_encoder(file[1])
            kwantyzacja = Kwant(uLawEncoded, bit)
            uLawDecoded = u_law_decoder(kwantyzacja)
            sd.play(uLawDecoded,file[2])
            sd.wait()
            input("Next song")

            print(f" {file[0]} kompresja DPCM bez perd, kwantyzacja {bit} bit")
            dpcmCompress = DPCM_compress(file[1], bit)
            dpcmDecompress = DPCM_decompress(dpcmCompress)
            sd.play(dpcmDecompress,file[2])
            sd.wait()
            input("Next song")

            print(f" {file[0]} kompresja DPCM z perd, kwantyzacja {bit} bit")
            dpcmCompress = DPCM_compress_pred(file[1], bit, predictor=predictor, n = 3)
            dpcmDecompress = DPCM_decompress_pred(dpcmCompress, predictor=predictor, n = 3)
            sd.play(dpcmDecompress,file[2])
            sd.wait()
            input("Next song")

document.save('lab6.docx')