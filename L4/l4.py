from docx import Document
from docx.shared import Inches

import numpy as np
import matplotlib.pyplot as plt
import sounddevice as sd
import soundfile as sf
import scipy.fftpack

from scipy.interpolate import interp1d
import matplotlib.pyplot as plt

from io import BytesIO

dataSin60, fsSin60 = sf.read('SIN/sin_60Hz.wav', dtype='int32')  
dataSin440, fsSin440 = sf.read('SIN/sin_440Hz.wav', dtype='int32')  
dataSin8000, fsSin8000 = sf.read('SIN/sin_8000Hz.wav', dtype='int32')  
dataSinComb, fsSinComb = sf.read('SIN/sin_combined.wav', dtype='int32')

dataSing1, fsSing1 = sf.read('SING/sing_high1.wav', dtype='float32') 
dataSing2, fsSing2 = sf.read('SING/sing_low2.wav', dtype='float32')  
dataSing3, fsSing3 = sf.read('SING/sing_medium2.wav', dtype='float32') 

Sin60 = ["sin_60Hz", dataSin60, fsSin60]
Sin440 = ["sin_440Hz",dataSin440, fsSin440]
Sin8000 = ["sin_8000Hz",dataSin8000, fsSin8000]
SinComb = ["sin_Comb",dataSinComb, fsSinComb]

Sing1 = ["sing_high1",dataSing1, fsSing1]
Sing2 = ["sing_low2",dataSing2, fsSing2]
Sing3 = ["sing_medium2",dataSing3, fsSing3]

frequencies = [2000, 4000, 8000, 11999, 16000, 16953, 24000, 41000]
bytes = [4,8,16,24]
decimations = [2,4,6,10,24]

freqSing = [4000 , 8000 ,11999 , 16000 , 16953]
bytesSing = [4, 8]
decimationsSing = [4,6,10,24]

generated = True

def Kwant(data,bit):
    d= 2 ** bit - 1
    if np.issubdtype(data.dtype,np.floating):
        m = -1
        n = 1
    else:
        n = np.iinfo(data.dtype).min
        m = np.iinfo(data.dtype).max
    DataF = data.astype(float)
    DataF = (DataF - m) / (n - m) * d
    DataF = np.round(DataF)
    DataF = DataF / d * (n - m) + m
    # kwantyzacja na DataF
    return DataF.astype(data.dtype)

def decimation(data, fs, n):
    return data[::n], fs/n

def interpolation(data, fs, nfs, type):
    N = data.shape[0]
    x=np.linspace(0,N,N)
    x1=np.linspace(0,N,int(N*nfs/fs))
    
    result = interp1d(x, data, kind=type)
    return result(x1).astype(data.dtype), nfs

def plotAudio(axs, Signal,Fs,TimeMargin=[0,0.02], fsize = 2**8):
    
    axs[0].plot(np.arange(0,Signal.shape[0])/Fs,Signal)
    axs[0].set(ylabel= "A", xlabel = "T [s]")
    axs[0].set(xlim = TimeMargin)

    yf = scipy.fftpack.fft(Signal,fsize)
    x = np.arange(0,Fs/2,Fs/fsize)
    y = 20*np.log10( np.abs(yf[:fsize//2]))
    axs[1].plot(x,y)
    axs[1].set(ylabel= "dB", xlabel = "F [Hz]")

SIN = [Sin60, Sin440, Sin8000, SinComb]
SING = [Sing1, Sing2, Sing3]
if(generated != True):
    document = Document()
    document.add_heading('Lab 4',0)
    for sin in SIN:
        document.add_heading('Plik - {}'.format(sin[0]),2)
        # initial
        fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
        plotAudio(axs, sin[1], sin[2])
        fig.suptitle('Initial') # Tytuł wykresu
        fig.tight_layout(pad=1.5) # poprawa czytelności 
        memfile = BytesIO() # tworzenie bufora
        fig.savefig(memfile) # z zapis do bufora 
        document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku        
        memfile.close()
        plt.close()

        document.add_heading(f'Kwantyzacja',2)
        for bit in bytes:
            document.add_heading(f'Kwantyzacja {sin[0]} bit: {bit}',3) # nagłówek sekcji, mozę być poziom wyżej
            fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
            data = Kwant(sin[1], bit)
            plotAudio(axs, data, sin[2])
            
            fig.suptitle(f'Kwantyzacja {sin[0]} bit:{bit}') # Tytuł wykresu
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora 
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku     
            memfile.close()
            plt.close()
            
        document.add_heading(f'Decymacja',2)
        for step in decimations:
            document.add_heading(f'Decymacja {sin[0]} step:{step}',3) # nagłówek sekcji, mozę być poziom wyżej
            fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
            data, fs = decimation(sin[1], sin[2], step)
            plotAudio(axs, data, fs)
            
            fig.suptitle(f'Decymacja {sin[0]} step:{step}' ) # Tytuł wykresu
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora 
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku        
            memfile.close()
            plt.close()
            
        document.add_heading(f'Interpolacja',2)
        for freq in frequencies:
            #liniowa
            document.add_heading(f'Interpolacja linear {sin[0]} freq: {freq}',3) # nagłówek sekcji, mozę być poziom wyżej
            fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
            data, fs = interpolation(sin[1], sin[2], freq, type='linear')
            plotAudio(axs, data, fs)
            
            fig.suptitle(f'Interpolacja linear {sin[0]} freq: {freq}') # Tytuł wykresu
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora 
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku        
            memfile.close()
            plt.close()
            #nieliniowa
            document.add_heading(f'Interpolacja nieliniowa {sin[0]} freq:{freq}',3) # nagłówek sekcji, mozę być poziom wyżej
            fig ,axs = plt.subplots(2,1,figsize=(10,7)) # tworzenie plota
            data, fs = interpolation(sin[1], sin[2], freq, type='cubic')
            plotAudio(axs, data, fs)
            
            fig.suptitle(f'Interpolacja nieliniowa {sin[0]} freq:{freq}') # Tytuł wykresu
            fig.tight_layout(pad=1.5) # poprawa czytelności 
            memfile = BytesIO() # tworzenie bufora
            fig.savefig(memfile) # z zapis do bufora 
            document.add_picture(memfile, width=Inches(6)) # dodanie obrazu z bufora do pliku
            memfile.close()
            plt.close()    
            
    document.save('lab4.docx') # zapis do plik

# AUDIO PLAYING
input("Press anything to start playing songs.")
for song in SING:
    print(f"oryginał {song[0]}")
    sd.play(song[1],song[2])
    sd.wait()
    input("Next song")
    
    for bit in bytesSing:
        print(f'Kwantyzacja {song[0]} {bit} bit')
        sd.play(Kwant(song[1], bit), song[2])
        sd.wait()
        input("Next song")

    for step in decimationsSing:
        print(f'Decymacja {song[0]} step {step}')
        data_dec, new_fs = decimation(song[1], song[2], step)
        sd.play(data_dec, new_fs)
        sd.wait()
        input("Next song")

    for newFs in freqSing:
        print(f'Interpolacja liniowa {song[0]} freq {newFs}')
        interp_data, new_fs = interpolation(song[1], song[2], newFs, type='linear')
        sd.play(interp_data, new_fs)
        sd.wait()
        input("Next song")

        print(f'Interpolacja nieliniowa {song[0]} freq {newFs}')
        interp_data, new_fs = interpolation(song[1], song[2], newFs, type='cubic')
        sd.play(interp_data, new_fs)
        sd.wait()
        input("Next song")