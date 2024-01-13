import numpy as np
import scipy.fftpack
import cv2
import matplotlib.pyplot as plt
import sys

from docx import Document
from docx.shared import Inches

def get_size(obj, seen=None):
    """Recursively finds size of objects"""
    size = sys.getsizeof(obj)
    if seen is None:
        seen = set()
    obj_id = id(obj)
    if obj_id in seen:
        return 0
    # Important mark as seen *before* entering recursion to gracefully handle
    # self-referential objects
    seen.add(obj_id)
    if isinstance(obj,np.ndarray):
        size=obj.nbytes
    elif isinstance(obj, dict):
        size += sum([get_size(v, seen) for v in obj.values()])
        size += sum([get_size(k, seen) for k in obj.keys()])
    elif hasattr(obj, '__dict__'):
        size += get_size(obj.__dict__, seen)
    elif hasattr(obj, '__iter__') and not isinstance(obj, (str, bytes, bytearray)):
        size += sum([get_size(i, seen) for i in obj])
    return size

class container:
    def __init__(self,Y,Cb,Cr,OGShape,Ratio="4:4:4",QY=np.ones((8,8)),QC=np.ones((8,8))):
        self.shape = OGShape
        self.Y=Y
        self.Cb=Cb
        self.Cr=Cr
        self.ChromaRatio=Ratio
        self.QY=QY
        self.QC=QC

QY= np.array([
        [16, 11, 10, 16, 24,  40,  51,  61],
        [12, 12, 14, 19, 26,  58,  60,  55],
        [14, 13, 16, 24, 40,  57,  69,  56],
        [14, 17, 22, 29, 51,  87,  80,  62],
        [18, 22, 37, 56, 68,  109, 103, 77],
        [24, 36, 55, 64, 81,  104, 113, 92],
        [49, 64, 78, 87, 103, 121, 120, 101],
        [72, 92, 95, 98, 112, 100, 103, 99],
        ])

QC= np.array([
        [17, 18, 24, 47, 99, 99, 99, 99],
        [18, 21, 26, 66, 99, 99, 99, 99],
        [24, 26, 56, 99, 99, 99, 99, 99],
        [47, 66, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        [99, 99, 99, 99, 99, 99, 99, 99],
        ])

QN= np.ones((8,8))

def rleEncoder(img):
    encoded = np.zeros(np.prod(img.shape)*2)
    img1d = img.flatten()
    info = np.array([len(img.shape)])
    info = np.concatenate([info, img.shape])
    i = 0
    bufferPosition = 0
    while i < len(img1d):
        counter = compareNext(img1d[i:len(img1d)])
        encoded[bufferPosition] = counter
        encoded[bufferPosition+1] =  img1d[i]
        bufferPosition += 2
        i += counter
    encoded = encoded[:bufferPosition]
    encoded = np.concatenate([info, encoded])
    return encoded.astype(int)

def rleDecoder(encoded):
    shape = encoded[1:int(encoded[0]+1)]
    decoded = []
    i = encoded[0] + 1
    while i < len(encoded):
        decoded = pushNumbers(decoded, encoded[i], encoded[i+1])    
        i += 2
    decoded = decoded.reshape(shape)
    return decoded

def compareNext(data):
    counter = 1
    for j in range(0, len(data) - 1):
        if(data[j] == data[j+1]):
            counter = counter + 1
        else:
            break
    return counter

def pushNumbers(array, count, number):
    array = np.array(array)
    array.resize(len(array) + int(count))
    array[-int(count):] = number
    return array

        
def CompressJPEG(RGB,Ratio="4:4:4",QY=np.ones((8,8)),QC=np.ones((8,8))):
    # RGB -> YCrCb
    YCrCb = cv2.cvtColor(RGB,cv2.COLOR_RGB2YCrCb).astype(int)
    # zapisać dane z wejścia do kalsy
    JPEG = container(Y = YCrCb[:,:,0], Cr=YCrCb[:,:,1], Cb=YCrCb[:,:,2], OGShape=RGB.shape, Ratio=Ratio, QY=QY, QC=QC)
    # Tu chroma subsampling

    JPEG.Cr = chroma_subsampling(JPEG.Cr, ratio=Ratio)
    JPEG.Cb = chroma_subsampling(JPEG.Cb, ratio=Ratio)
    
    JPEG.Y=CompressLayer(JPEG.Y,JPEG.QY)
    JPEG.Cr=CompressLayer(JPEG.Cr,JPEG.QC)
    JPEG.Cb=CompressLayer(JPEG.Cb,JPEG.QC)

    # tu dochodzi kompresja bezstratna
    JPEG.Y = rleEncoder(JPEG.Y)
    JPEG.Cr = rleEncoder(JPEG.Cr)
    JPEG.Cb = rleEncoder(JPEG.Cb)
    
    yCompression = get_size(JPEG.Y) 
    CrCompression = get_size(JPEG.Cr) 
    CbCompression = get_size(JPEG.Cb)
    
    compressions = [yCompression, CrCompression, CbCompression]
    
    return JPEG, compressions

def DecompressJPEG(JPEG, Ratio="4:4:4"):
    
    # dekompresja bezstratna
    JPEG.Y = rleDecoder(JPEG.Y)
    JPEG.Cr = rleDecoder(JPEG.Cr)
    JPEG.Cb = rleDecoder(JPEG.Cb)
    
    Y=DecompressLayer(JPEG.Y,JPEG.QY, JPEG.shape)
    Cr=DecompressLayer(JPEG.Cr,JPEG.QC, JPEG.shape, ratio=Ratio)
    Cb=DecompressLayer(JPEG.Cb,JPEG.QC, JPEG.shape, ratio=Ratio)
    # Tu chroma resampling
    Cr = chroma_resampling(Cr, ratio=Ratio)
    Cb = chroma_resampling(Cb, ratio=Ratio)
    # tu rekonstrukcja obrazus
    # YCrCb -> RGB
    YCrCb=np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)
    RGB=cv2.cvtColor(YCrCb.astype(np.uint8),cv2.COLOR_YCrCb2RGB)
    return RGB

def CompressBlock(block,Q):
    #dtc
    block = dct2(block)
    
    #kwant
    block = np.round(block / Q).astype(int)
    
    #zigzag
    vector = zigzag(block)
    
    return vector

def DecompressBlock(vector,Q):
    #zigzag
    block = zigzag(vector)
    
    #kwant
    block = block * Q
    
    #dtc
    block = idct2(block)
    return block


## podział na bloki
# L - warstwa kompresowana
# S - wektor wyjściowy
def CompressLayer(L,Q):
    S=np.array([])
    for w in range(0,L.shape[0],8):
        for k in range(0,L.shape[1],8):
            block=L[w:(w+8),k:(k+8)]
            S=np.append(S, CompressBlock(block,Q))
    return S

## wyodrębnianie bloków z wektora 
# L - warstwa o oczekiwanym rozmiarze
# S - długi wektor zawierający skompresowane dane
def DecompressLayer(S, Q, OGShape, ratio = "4:4:4"):
    #zadeklaruj odpowiedniego rozmiaru macierzy
    if(ratio == "4:2:2"):
        L = np.zeros([OGShape[0], int(OGShape[1]) // 2])
    elif(ratio == "4:2:0"):
        L = np.zeros([int(OGShape[0] // 2), int(OGShape[1] // 2)])
    else: 
        L = np.zeros([OGShape[0], OGShape[1]])
    for idx,i in enumerate(range(0,S.shape[0],64)):
        vector=S[i:(i+64)]
        m=L.shape[1]/8
        k=int((idx%m)*8)
        w=int((idx//m)*8)
        block = DecompressBlock(vector,Q)
        L[w:(w+8),k:(k+8)] = block
    return L
        
def zigzag(A):
    template = np.array([
            [0,  1,  5,  6,  14, 15, 27, 28],
            [2,  4,  7,  13, 16, 26, 29, 42],
            [3,  8,  12, 17, 25, 30, 41, 43],
            [9,  11, 18, 24, 31, 40, 44, 53],
            [10, 19, 23, 32, 39, 45, 52, 54],
            [20, 22, 33, 38, 46, 51, 55, 60],
            [21, 34, 37, 47, 50, 56, 59, 61],
            [35, 36, 48, 49, 57, 58, 62, 63],
            ])
    if len(A.shape)==1:
        B=np.zeros((8,8))
        for r in range(0,8):
            for c in range(0,8):
                B[r,c]=A[template[r,c]]
    else:
        B=np.zeros((64,))
        for r in range(0,8):
            for c in range(0,8):
                B[template[r,c]]=A[r,c]
    return B        

def dct2(a):
    return scipy.fftpack.dct( scipy.fftpack.dct( a.astype(float), axis=0, norm='ortho' ), axis=1, norm='ortho' )

def idct2(a):
    return scipy.fftpack.idct( scipy.fftpack.idct( a.astype(float), axis=0 , norm='ortho'), axis=1 , norm='ortho')

def chroma_subsampling(img, ratio):
    if ratio=="4:2:2":
        new_img = img[:,::2]
    elif ratio=="4:2:0":
        new_img = img[::2,::2]
    else:
        new_img = img
    return new_img

def chroma_resampling(img, ratio):
    if ratio == "4:2:2":
        resampled = np.repeat(img, repeats=2, axis=1)
    elif ratio == "4:2:0":
        resampled = np.repeat(np.repeat(img, repeats=2, axis=0), repeats=2, axis=1)
    else:
        resampled = img
    return resampled


files = [ "imgs/mountain.png", "imgs/cars.png", "imgs/food.png", "imgs/capybara.png"]
titles = ["mountain", "cars", "food", "capybara"]
sections = [[200,328, 200,328], [400,528,400,528], [130,258,130,258]]
ratios = ["4:4:4", "4:2:2", "4:2:0"]

document = Document()
document.add_heading('Lab 7',0)

for picTitle, file in enumerate(files):
    print(f"current file: {file}")
    document.add_heading(f'{file}', 1)
    img = cv2.cvtColor(cv2.imread(file), cv2.COLOR_BGR2RGB)
    for i, section in enumerate(sections):
        for j, ratio in enumerate(ratios):
            document.add_heading(f'{file}, section {section}, ratio {ratio}', 2)
            imgPart = img[section[0]:section[1], section[2]:section[3]]
            fig, axs = plt.subplots(4, 2, sharey=True)
            fig.set_size_inches(9,13)
            # obraz oryginalny
            YCrCb= cv2.cvtColor(imgPart,cv2.COLOR_RGB2YCrCb).astype(int)
            axs[0,0].imshow(imgPart) #RGB 
            # jego warstwy w Y Cr Cb -> dopisać konwersję
            axs[1,0].imshow(YCrCb[:,:,0],cmap=plt.cm.gray) 
            axs[2,0].imshow(YCrCb[:,:,1],cmap=plt.cm.gray)
            axs[3,0].imshow(YCrCb[:,:,2],cmap=plt.cm.gray)

            compressed, compressions = CompressJPEG(imgPart,ratio)

            yCompression = compressions[0] / (128*128) * 100
            CrCompression = compressions[1] / (128*128) * 100
            CbCompression = compressions[2] / (128*128) * 100
            document.add_paragraph(f"Y compression: {yCompression} % , Cr compression {CrCompression}% , Cb compression {CbCompression}%")

            decompressed = DecompressJPEG(compressed, ratio)

            # obraz po dekompresji
            axs[0,1].imshow(decompressed) #RGB 
            YCrCbDec= cv2.cvtColor(decompressed,cv2.COLOR_RGB2YCrCb).astype(int)
            # jego warstwy w Y Cr Cb -> dopisać konwersję
            axs[1,1].imshow(YCrCbDec[:,:,0],cmap=plt.cm.gray)
            axs[2,1].imshow(YCrCbDec[:,:,1],cmap=plt.cm.gray)
            axs[3,1].imshow(YCrCbDec[:,:,2],cmap=plt.cm.gray)
            
            fig.savefig(f"figs/{titles[picTitle]}{i}{j}.png")
            document.add_picture(f"figs/{titles[picTitle]}{i}{j}.png", width=Inches(5), height=Inches(6))
        
document.save("l7.docx")