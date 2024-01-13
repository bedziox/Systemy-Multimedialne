import cv2
import numpy as np
import matplotlib.pyplot as plt
import os


from docx import Document
from docx.shared import Inches

##############################################################################
######   Konfiguracja       ##################################################
##############################################################################

kat='.\wideo\\'                                             # katalog z plikami wideo
plik = "clip_1.mp4"                                         # plik
ile= 100                                                    # ile klatek odtworzyć? <0 - całość
key_frame_counter=4                                         # co która klatka ma być kluczowa i nie podlegać kompresji
plot_frames=np.array([30,45])                               # automatycznie wyrysuj wykresy
auto_pause_frames=np.array([26])                            # automatycznie za pauzuj dla klatki
subsamplings=["4:4:4", "4:2:2", "4:4:0", "4:1:1", "4:1:0"]  # parametry dla chroma subsampling
dzielnik= 0.25                                                  # dzielnik przy zapisie różnicy
dzielniki = [1, 0.5, 0.25]                                    # dzieniki
wyswietlaj_klatki=False                                     # czy program ma wyświetlać klatki
ROI = [[70,170,460,560], [345, 445, 1100, 1200]]            # wyświetlane fragmenty (można podać kilka )



def rleEncoder(img):
    encoded = np.zeros(np.prod(img.shape)*2)
    img1d = img.flatten()
    info = np.array([len(img.shape)])
    info = np.concatenate([info, img.shape])
    i = 0
    bufferPosition = 0
    while i < len(img1d):
        counter = compareNext(img1d[i:len(img1d)])
        encoded[bufferPosition] = counter
        encoded[bufferPosition+1] =  img1d[i]
        bufferPosition += 2
        i += counter
    encoded = encoded[:bufferPosition]
    encoded = np.concatenate([info, encoded])
    return encoded.astype(int)

def rleDecoder(encoded):
    shape = encoded[1:int(encoded[0]+1)]
    decoded = []
    i = encoded[0] + 1
    while i < len(encoded):
        decoded = pushNumbers(decoded, encoded[i], encoded[i+1])    
        i += 2
    decoded = decoded.reshape(shape)
    return decoded

def compareNext(data):
    counter = 1
    for j in range(0, len(data) - 1):
        if(data[j] == data[j+1]):
            counter = counter + 1
        else:
            break
    return counter

def pushNumbers(array, count, number):
    array = np.array(array)
    array.resize(len(array) + int(count))
    array[-int(count):] = number
    return array


##############################################################################
####     Kompresja i dekompresja    ##########################################
##############################################################################
class data:
    def init(self):
        self.Y=None
        self.Cb=None
        self.Cr=None

def Chroma_subsampling(L,subsampling):
    if subsampling == "4:2:2":
        L = L[::,::2]
    elif subsampling == "4:2:0":
        L = L[::2, ::2]
    elif subsampling == "4:4:0":
        L = L[::2,::]
    elif subsampling == "4:1:1":
        L = L[::,::4]
    elif subsampling == "4:1:0":
        L = L[::2,::4]
    else:
        L = L
    return L

def Chroma_resampling(L,subsampling):
    if subsampling == "4:2:2":
        L = np.repeat(L, repeats=2, axis=1)
    elif subsampling == "4:2:0":
        L = np.repeat(np.repeat(L, repeats=2, axis=0), repeats=2, axis=1)
    elif subsampling == "4:4:0":
        L = np.repeat(L, repeats=2, axis=0)
    elif subsampling == "4:1:1":
        L = np.repeat(L, repeats=4, axis=1)
    elif subsampling == "4:1:0":
        L = np.repeat(np.repeat(L, repeats=2, axis=0), repeats=4, axis=1)
    else:
        L = L
    return L

        
def frame_image_to_class(frame,subsampling):
    Frame_class = data()
    Frame_class.Y=frame[:,:,0].astype(int)
    Frame_class.Cb=Chroma_subsampling(frame[:,:,2].astype(int),subsampling)
    Frame_class.Cr=Chroma_subsampling(frame[:,:,1].astype(int),subsampling)
    return Frame_class


def frame_layers_to_image(Y,Cr,Cb,subsampling):  
    Cb=Chroma_resampling(Cb,subsampling)
    Cr=Chroma_resampling(Cr,subsampling)
    return np.dstack([Y,Cr,Cb]).clip(0,255).astype(np.uint8)

def compress_KeyFrame(Frame_class):
    KeyFrame = data()
    ## TO DO 
    KeyFrame.Y=Frame_class.Y
    KeyFrame.Cb=Frame_class.Cb
    KeyFrame.Cr=Frame_class.Cr
    
    # KeyFrame.Y=rleEncoder(Frame_class.Y)
    # KeyFrame.Cb=rleEncoder(Frame_class.Cb)
    # KeyFrame.Cr=rleEncoder(Frame_class.Cr)
    return KeyFrame

def decompress_KeyFrame(KeyFrame):
    # Y=rleDecoder(KeyFrame.Y)
    # Cb=rleDecoder(KeyFrame.Cb)
    # Cr=rleDecoder(KeyFrame.Cr)
    
    Y=KeyFrame.Y
    Cb=KeyFrame.Cb
    Cr=KeyFrame.Cr
    ## TO DO 
    frame_image=frame_layers_to_image(Y,Cr,Cb,subsampling)
    
    return frame_image

def compress_not_KeyFrame(Frame_class, KeyFrame, Rnum = 1, inne_paramerty_do_dopisania=None):
    Compress_data = data()
    ## TO DO 
    Compress_data.Y = (Frame_class.Y - KeyFrame.Y) * Rnum
    Compress_data.Cb = (Frame_class.Cb - KeyFrame.Cb) * Rnum
    Compress_data.Cr = (Frame_class.Cr - KeyFrame.Cr) * Rnum
    
    # Compress_data.Y=rleEncoder(Compress_data.Y)
    # Compress_data.Cb=rleEncoder(Compress_data.Cb)
    # Compress_data.Cr=rleEncoder(Compress_data.Cr)
    return Compress_data

def decompress_not_KeyFrame(Compress_data,  KeyFrame , Rnum = 1, inne_paramerty_do_dopisania=None):
    
    Y = KeyFrame.Y + Compress_data.Y * 1/Rnum
    Cb = KeyFrame.Cb + Compress_data.Cb * 1/Rnum
    Cr = KeyFrame.Cr + Compress_data.Cr * 1/Rnum
    
    # Y=rleDecoder(Compress_data.Y)
    # Cb=rleDecoder(Compress_data.Cb)
    # Cr=rleDecoder(Compress_data.Cr)
    ## TO DO
    return frame_layers_to_image(Y,Cr,Cb,subsampling)

def plotDiffrence(ReferenceFrame,DecompressedFrame,ROI,sampl,R):
    # bardzo słaby i sztuczny przykład wykorzystania tej opcji
    # przerobić żeby porównanie było dokonywane w RGB nie YCrCb i/lub zastąpić innym porównaniem
    # ROI - Region of Insert współrzędne fragmentu który chcemy przybliżyć i ocenić w formacie [w1,w2,k1,k2]
    fig, axs = plt.subplots(1, 3 , sharey=True   )
    fig.set_size_inches(16,5)
    
    ReferenceFrame = cv2.cvtColor(ReferenceFrame, cv2.COLOR_YCrCb2RGB)
    DecompressedFrame = cv2.cvtColor(DecompressedFrame, cv2.COLOR_YCrCb2RGB)
    
    axs[0].imshow(ReferenceFrame[ROI[0]:ROI[1],ROI[2]:ROI[3]])
    axs[2].imshow(DecompressedFrame[ROI[0]:ROI[1],ROI[2]:ROI[3]]) 
    diff=ReferenceFrame[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)-DecompressedFrame[ROI[0]:ROI[1],ROI[2]:ROI[3]].astype(float)
    print(np.min(diff),np.max(diff))
    axs[1].imshow(diff,vmin=np.min(diff),vmax=np.max(diff))
    fig.suptitle(f"{plik.split('.')[0]} roi {ROI} sampl {sampl} r {R}diff.png")
    fig.savefig(f"diffs/{plik.split('.')[0]}roi{ROI}sampl{sampl}r{R}diff.png")
    document.add_picture(f"diffs/{plik.split('.')[0]}roi{ROI}sampl{sampl}r{R}diff.png", width=Inches(5))
    plt.close()


##############################################################################
####     Głowna pętla programu      ##########################################
##############################################################################

generate = True
document = Document()
document.add_heading('Lab 8',0)

for dzielnik in dzielniki:
    document.add_heading("dzielnik {}".format(dzielnik))
    for index, subsampling in enumerate(subsamplings):
        document.add_heading("subsampling {}".format(subsampling))
        compression_information=np.zeros((3,ile))
        cap = cv2.VideoCapture(kat+'\\'+plik)
        if ile<0:
            ile=int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        for i in range(ile):
            ret, frame = cap.read()
            if wyswietlaj_klatki:
                cv2.namedWindow('Normal Frame')
                cv2.imshow('Normal Frame',frame)
            frame=cv2.cvtColor(frame,cv2.COLOR_BGR2YCrCb)
            Frame_class = frame_image_to_class(frame,subsampling)
            if (i % key_frame_counter)==0: # pobieranie klatek kluczowych
                KeyFrame = compress_KeyFrame(Frame_class)
                cY=KeyFrame.Y
                cCb=KeyFrame.Cb
                cCr=KeyFrame.Cr
                Decompresed_Frame = decompress_KeyFrame(KeyFrame)
            else: # kompresja
                Compress_data = compress_not_KeyFrame(Frame_class, KeyFrame)
                cY=Compress_data.Y
                cCb=Compress_data.Cb
                cCr=Compress_data.Cr
                Decompresed_Frame = decompress_not_KeyFrame(Compress_data,  KeyFrame)
            
            compression_information[0,i]= (frame[:,:,0].size - cY.size)/frame[:,:,0].size
            compression_information[1,i]= (frame[:,:,0].size - cCb.size)/frame[:,:,0].size
            compression_information[2,i]= (frame[:,:,0].size - cCr.size)/frame[:,:,0].size  
            if wyswietlaj_klatki:
                cv2.namedWindow('Decompressed Frame')
                cv2.imshow('Decompressed Frame',cv2.cvtColor(Decompresed_Frame,cv2.COLOR_YCrCb2BGR))
            
            if np.any(plot_frames==i): # rysuj wykresy
                for r in ROI:
                    plotDiffrence(frame,Decompresed_Frame,r,subsampling.replace(':',''),dzielnik)
                
            if np.any(auto_pause_frames==i):
                cv2.waitKey(-1) #wait until any key is pressed
            
            k = cv2.waitKey(1) & 0xff
            
            if k==ord('q'):
                break
            elif k == ord('p'):
                cv2.waitKey(-1) #wait until any key is pressed
        
        plt.figure()
        plt.plot(np.arange(0,ile),compression_information[0,:]*100)
        plt.plot(np.arange(0,ile),compression_information[1,:]*100)
        plt.plot(np.arange(0,ile),compression_information[2,:]*100)
        plt.ylim(0, 100)
        plt.legend(['Y', 'Cr', 'Cb'])
        plt.title("File:{}, subsampling={}, divider={}, KeyFrame={}".format(plik,subsampling,dzielnik,key_frame_counter))
        plt.savefig(f"imgs/{plik.split('.')[0]}div{dzielnik}sampl{subsampling.replace(':','')}.png")
        document.add_heading(f"imgs/{plik.split('.')[0]}div{dzielnik}sampl{subsampling.replace(':','')}.png", 2)
        document.add_picture(f"imgs/{plik.split('.')[0]}div{dzielnik}sampl{subsampling.replace(':','')}.png", width=Inches(5))
        plt.close()

document.save("Lab8.docx")